"""
@file docx_service.py
@brief Word 文档生成模块。

该模块优先使用 docxtpl 渲染基础模板，再使用 python-docx 动态追加
文档结构树中的多级标题、正文和表格。

@author TextTreeDoc 项目组
@date 2026
"""

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
from docxtpl import DocxTemplate

from app.core.config import GENERATED_DIR, REPORT_TEMPLATE_PATH, ensure_runtime_dirs
from app.core.database import get_connection
from app.services.image_service import get_image

TEMPLATE_VERSION = "TextTreeDoc Template v4"
BODY_FONT = "宋体"
HEADING_FONT = "黑体"
ASCII_FONT = "Times New Roman"
TEXT_COLOR = RGBColor(0, 0, 0)
DEFAULT_FORMAT_CONFIG: dict[str, Any] = {
    "heading_numbering": "decimal",
    "cover": False,
    "cover_style": "none",
    "show_title": True,
    "body_font": BODY_FONT,
    "heading_font": HEADING_FONT,
    "ascii_font": ASCII_FONT,
    "body_size": 12,
    "heading1_size": 16,
    "heading2_size": 14,
    "heading3_size": 12,
    "line_spacing": 1.5,
    "line_spacing_rule": {"type": "multiple", "value": 1.5, "unit": "line"},
    "first_line_indent_chars": 2,
    "paragraph_space_after": 6,
    "body_space_before": {"value": 0, "unit": "pt"},
    "body_space_after": {"value": 6, "unit": "pt"},
    "heading1_space_before": {"value": 18, "unit": "pt"},
    "heading1_space_after": {"value": 12, "unit": "pt"},
    "heading2_space_before": {"value": 12, "unit": "pt"},
    "heading2_space_after": {"value": 6, "unit": "pt"},
    "heading3_space_before": {"value": 6, "unit": "pt"},
    "heading3_space_after": {"value": 6, "unit": "pt"},
    "table_style": "Table Grid",
}


def ensure_report_template() -> None:
    """
    @brief 确保 docxtpl 报告模板存在。

    @return None。

    没有模板文件时自动创建最小模板，避免项目首次运行时无法生成 Word。
    """
    ensure_runtime_dirs()
    if REPORT_TEMPLATE_PATH.exists() and _template_is_current():
        return
    document = Document()
    _configure_document_styles(document)
    document.core_properties.comments = TEMPLATE_VERSION
    document.add_paragraph("{{ title }}")
    document.save(REPORT_TEMPLATE_PATH)


def generate_docx_from_tree(title: str, tree: dict[str, Any], format_config: dict[str, Any] | None = None) -> dict:
    """
    @brief 根据文档结构树生成 docx 文件。

    @param title 文档标题。
    @param tree 文档结构树，包含 title、sections、children 和 blocks 字段。
    @return 包含 filename、path 和 download_url 的字典。
    """
    config = _merge_format_config(format_config)
    ensure_report_template()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"report_{timestamp}.docx"
    output_path = GENERATED_DIR / filename

    template = DocxTemplate(REPORT_TEMPLATE_PATH)
    template.render({"title": title, "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")})
    template.save(output_path)

    document = Document(output_path)
    _clear_document_body(document)
    _configure_document_styles(document, config)
    front_matter, body_sections, back_matter = _split_document_sections(tree.get("sections", []))
    _append_front_matter(document, title, front_matter, config)
    for index, section in enumerate(body_sections, 1):
        _append_section(document, section, level=1, config=config, path=[index])
    _append_back_matter(document, back_matter, config)
    document.save(output_path)
    _save_document_record(title, tree, output_path)
    return {
        "filename": filename,
        "path": str(output_path),
        "download_url": f"/documents/download/{filename}",
    }


def _append_section(
    document: Document,
    section: dict[str, Any],
    level: int,
    config: dict[str, Any],
    path: list[int],
) -> None:
    """
    @brief 递归追加文档章节。

    @param document python-docx 文档对象。
    @param section 单个章节节点。
    @param level 当前标题级别。
    @return None。
    """
    heading = section.get("heading")
    if heading:
        heading_size = _heading_size(level, config)
        paragraph = document.add_heading(_format_heading(heading, level, path, config), level=min(level, 4))
        _apply_paragraph_font(
            paragraph,
            config["heading_font"],
            config["ascii_font"],
            size=heading_size,
            bold=True,
            space_before=_heading_spacing(level, config, "before", heading_size),
            space_after=_heading_spacing(level, config, "after", heading_size),
        )
    paragraphs = section.get("paragraphs") or []
    if not paragraphs and section.get("content"):
        paragraphs = [section.get("content")]
    for content in paragraphs:
        if not content:
            continue
        paragraph = document.add_paragraph(str(content))
        _apply_paragraph_font(
            paragraph,
            config["body_font"],
            config["ascii_font"],
            size=config["body_size"],
            first_line_indent_chars=config["first_line_indent_chars"],
            line_spacing=_resolve_line_spacing(config),
            space_before=_resolve_spacing(config["body_space_before"], config["body_size"]),
            space_after=_resolve_spacing(config["body_space_after"], config["body_size"]),
        )
    for block in section.get("blocks", []):
        _append_block(document, block, config)
    for index, child in enumerate(section.get("children", []), 1):
        _append_section(document, child, level=level + 1, config=config, path=[*path, index])


def _clear_document_body(document: Document) -> None:
    """
    @brief 清空 docxtpl 基础模板正文。

    @param document Word 文档对象。
    @return None。
    """
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _split_document_sections(sections: list[dict[str, Any]]) -> tuple[dict[str, dict[str, Any]], list[dict[str, Any]], dict[str, list[dict[str, Any]]]]:
    """
    @brief 将结构树章节拆分为前置页、正文和后置页。

    @param sections 原始章节列表。
    @return front_matter、body_sections、back_matter。
    """
    front_matter: dict[str, dict[str, Any]] = {}
    back_matter: dict[str, list[dict[str, Any]]] = {"references": [], "appendix": [], "teacher_comment": []}
    body_sections: list[dict[str, Any]] = []
    body_container: dict[str, Any] | None = None
    for section in sections:
        kind = _section_kind(section)
        if kind in {"cover", "declaration", "abstract", "toc"}:
            front_matter[kind] = section
        elif kind in back_matter:
            back_matter[kind].append(section)
        elif kind == "body":
            body_container = section
        else:
            body_sections.append(section)
    if body_container:
        children = body_container.get("children") or []
        if children:
            body_sections = [*children, *body_sections]
        elif body_container.get("paragraphs") or body_container.get("content"):
            body_sections.insert(0, body_container)
    return front_matter, body_sections, back_matter


def _section_kind(section: dict[str, Any]) -> str:
    """
    @brief 判断章节类型。

    @param section 章节。
    @return 类型标识。
    """
    heading = _strip_heading_prefix(str(section.get("heading") or "")).replace(" ", "")
    explicit_type = str(section.get("type") or "").lower()
    if explicit_type:
        return explicit_type
    if "封面" in heading:
        return "cover"
    if "声明" in heading:
        return "declaration"
    if "摘要" in heading:
        return "abstract"
    if heading == "目录" or "目录" in heading:
        return "toc"
    if heading == "正文":
        return "body"
    if "参考文献" in heading:
        return "references"
    if "附录" in heading:
        return "appendix"
    if "教师评语" in heading or "评分" in heading:
        return "teacher_comment"
    return "body_section"


def _append_front_matter(document: Document, title: str, front_matter: dict[str, dict[str, Any]], config: dict[str, Any]) -> None:
    """
    @brief 渲染封面、声明、摘要和目录。

    @param document Word 文档对象。
    @param title 文档标题。
    @param front_matter 前置页数据。
    @param config 格式配置。
    @return None。
    """
    has_cover = _should_render_cover(front_matter, config)
    if has_cover:
        _append_cover_page(document, title, front_matter.get("cover"), config)
    elif config.get("show_title", True):
        _append_document_title(document, title, config)
    if front_matter.get("declaration"):
        _append_declaration_page(document, front_matter["declaration"], config)
    if front_matter.get("abstract") or config.get("abstract"):
        _append_abstract_page(document, front_matter.get("abstract"), config)
    if front_matter.get("toc") or config.get("toc"):
        _append_toc_placeholder(document, config)


def _append_document_title(document: Document, title: str, config: dict[str, Any]) -> None:
    """
    @brief 在无封面文档开头追加正文标题。

    @param document Word 文档对象。
    @param title 文档标题。
    @param config 格式配置。
    @return None。
    """
    title_paragraph = document.add_paragraph(str(title or "未命名文档"))
    title_paragraph.alignment = 1
    _apply_paragraph_font(
        title_paragraph,
        config["heading_font"],
        config["ascii_font"],
        size=max(_as_float(config.get("heading1_size"), 16) + 2, 18),
        bold=True,
        line_spacing=_resolve_line_spacing(config),
        space_before=Pt(0),
        space_after=Pt(18),
    )


def _append_back_matter(document: Document, back_matter: dict[str, list[dict[str, Any]]], config: dict[str, Any]) -> None:
    """
    @brief 渲染参考文献、附录和教师评语等后置页。

    @param document Word 文档对象。
    @param back_matter 后置页数据。
    @param config 格式配置。
    @return None。
    """
    for key in ("references", "appendix", "teacher_comment"):
        for section in back_matter.get(key, []):
            document.add_page_break()
            _append_special_section(document, section, config)


def _append_cover_page(document: Document, title: str, section: dict[str, Any] | None, config: dict[str, Any]) -> None:
    """
    @brief 渲染通用封面页。

    @param document Word 文档对象。
    @param title 文档标题。
    @param section 封面结构节点。
    @param config 格式配置。
    @return None。
    """
    school = document.add_paragraph("武汉大学计算机学院")
    school.alignment = 1
    _apply_paragraph_font(school, "宋体", config["ascii_font"], size=26)
    report = document.add_paragraph("本科生课程设计报告")
    report.alignment = 1
    _apply_paragraph_font(report, "宋体", config["ascii_font"], size=26)
    document.add_paragraph()
    title_paragraph = document.add_paragraph(title)
    title_paragraph.alignment = 1
    _apply_paragraph_font(title_paragraph, "黑体", config["ascii_font"], size=22, bold=True, line_spacing=Pt(32))
    document.add_paragraph()
    fields = _extract_cover_fields(section)
    for label, value in fields:
        paragraph = document.add_paragraph(f"{label:<8}：{value}")
        paragraph.alignment = 1
        _apply_paragraph_font(paragraph, "宋体", config["ascii_font"], size=15)
    document.add_paragraph()
    date_paragraph = document.add_paragraph(datetime.now().strftime("%Y年%m月"))
    date_paragraph.alignment = 1
    _apply_paragraph_font(date_paragraph, "宋体", config["ascii_font"], size=15)
    document.add_page_break()


def _should_render_cover(front_matter: dict[str, dict[str, Any]], config: dict[str, Any]) -> bool:
    """
    @brief 判断是否渲染封面。

    @param front_matter 前置页配置。
    @param config 格式配置。
    @return 是否渲染封面。
    """
    if not (front_matter.get("cover") or config.get("cover")):
        return False
    return config.get("cover_style") == "wuhan_cs_course_design" or bool(front_matter.get("cover"))


def _append_declaration_page(document: Document, section: dict[str, Any], config: dict[str, Any]) -> None:
    title = document.add_paragraph("郑 重 声 明")
    title.alignment = 1
    _apply_paragraph_font(title, "宋体", config["ascii_font"], size=22, bold=True)
    for text in _section_texts(section) or [
        "本人呈交的设计报告，是在指导老师的指导下，独立进行实验工作所取得的成果，所有数据、图片资料真实可靠。除文中已经注明引用的内容外，本设计报告不包含他人享有著作权的内容。"
    ]:
        paragraph = document.add_paragraph(text)
        _apply_paragraph_font(paragraph, "宋体", config["ascii_font"], size=14, first_line_indent_chars=2, line_spacing=_resolve_line_spacing(config))
    sign = document.add_paragraph("本人签名：                     日期：")
    _apply_paragraph_font(sign, "宋体", config["ascii_font"], size=14)
    document.add_page_break()


def _append_abstract_page(document: Document, section: dict[str, Any] | None, config: dict[str, Any]) -> None:
    title = document.add_paragraph("摘  要")
    title.alignment = 1
    _apply_paragraph_font(title, "黑体", config["ascii_font"], size=18, bold=True)
    texts = _section_texts(section) if section else []
    for text in texts or ["本文围绕课程设计主题，对系统需求、设计过程、实现方法和结果进行概括说明。"]:
        paragraph = document.add_paragraph(text)
        _apply_paragraph_font(paragraph, config["body_font"], config["ascii_font"], size=config["body_size"], first_line_indent_chars=2, line_spacing=_resolve_line_spacing(config))
    keywords = document.add_paragraph("关键词：")
    _apply_paragraph_font(keywords, "黑体", config["ascii_font"], size=12, bold=True)
    document.add_page_break()


def _append_toc_placeholder(document: Document, config: dict[str, Any]) -> None:
    title = document.add_paragraph("目  录")
    title.alignment = 1
    _apply_paragraph_font(title, "黑体", config["ascii_font"], size=18, bold=True)
    _append_toc_field(document)
    document.add_page_break()


def _append_toc_field(document: Document) -> None:
    """
    @brief 插入 Word 目录域。

    @param document Word 文档对象。
    @return None。
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fallback_text = OxmlElement("w:t")
    fallback_text.text = "请在 Word 中右键更新目录"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, fallback_text, fld_char_end])


def _append_special_section(document: Document, section: dict[str, Any], config: dict[str, Any]) -> None:
    heading = _strip_heading_prefix(section.get("heading") or "参考文献")
    paragraph = document.add_heading(heading, level=1)
    _apply_paragraph_font(paragraph, config["heading_font"], config["ascii_font"], size=config["heading1_size"], bold=True)
    for text in _section_texts(section):
        body = document.add_paragraph(text)
        _apply_paragraph_font(body, config["body_font"], config["ascii_font"], size=config["body_size"], line_spacing=_resolve_line_spacing(config))


def _section_texts(section: dict[str, Any] | None) -> list[str]:
    if not section:
        return []
    texts = [str(text) for text in section.get("paragraphs") or [] if text]
    if not texts and section.get("content"):
        texts.append(str(section["content"]))
    return texts


def _extract_cover_fields(section: dict[str, Any] | None) -> list[tuple[str, str]]:
    content = " ".join(_section_texts(section))
    defaults = [
        ("专业名称", "XXX XXX"),
        ("课程名称", "XXX XXX"),
        ("指导教师一", "XXX    职称"),
        ("指导教师二", "XXX    职称"),
        ("学生学号", "20XXXXXXXXX"),
        ("学生姓名", "XXX"),
    ]
    return [(label, _extract_field_value(content, label) or default) for label, default in defaults]


def _extract_field_value(content: str, label: str) -> str:
    import re

    compact_label = r"\s*".join(label)
    stop_labels = "专业名称|课程名称|指导教师一|指导教师二|学生学号|学生姓名"
    match = re.search(
        rf"{compact_label}\s*[：:]\s*(.*?)(?=\s*(?:{stop_labels})\s*[：:]|$)",
        content,
    )
    return match.group(1).strip() if match else ""


def _append_block(document: Document, block: dict[str, Any], config: dict[str, Any]) -> None:
    """
    @brief 将结构树 block 写入 Word。

    @param document python-docx 文档对象。
    @param block 支持 table 类型，预留 image 类型。
    @return None。
    """
    if block.get("type") == "image":
        _append_image_block(document, block, config)
        return
    if block.get("type") == "table":
        headers = block.get("headers", [])
        rows = block.get("rows", [])
        if not headers:
            return
        table = document.add_table(rows=1, cols=len(headers))
        try:
            table.style = config["table_style"]
        except (KeyError, ValueError):
            table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = str(header)
            _apply_cell_font(table.rows[0].cells[index], config, bold=True)
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row[: len(headers)]):
                cells[index].text = str(value)
                _apply_cell_font(cells[index], config)


def _append_image_block(document: Document, block: dict[str, Any], config: dict[str, Any]) -> None:
    """
    @brief 将 image 类型 block 写入 Word。

    @param document python-docx 文档对象。
    @param block 图片块，包含 image_id、caption 等字段。
    @param config 文档格式配置。
    @return None。
    """
    image_id = block.get("image_id")
    if not image_id:
        return
    try:
        image = get_image(int(image_id))
    except (TypeError, ValueError):
        return
    image_path = Path(image["file_path"])
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(5.4))
    except Exception:
        return
    caption = block.get("caption") or image.get("caption") or image.get("name")
    if caption:
        caption_paragraph = document.add_paragraph(str(caption))
        caption_paragraph.alignment = 1
        _apply_paragraph_font(
            caption_paragraph,
            config["body_font"],
            config["ascii_font"],
            size=max(9, config["body_size"] - 1),
            line_spacing=1.0,
            space_before=Pt(3),
            space_after=Pt(6),
        )


def _template_is_current() -> bool:
    """
    @brief 判断当前模板是否为最新版本。

    @return 模板版本匹配时返回 True。
    """
    try:
        document = Document(REPORT_TEMPLATE_PATH)
    except Exception:
        return False
    return document.core_properties.comments == TEMPLATE_VERSION


def _configure_document_styles(document: Document, config: dict[str, Any] | None = None) -> None:
    """
    @brief 配置 Word 文档的中文正文和标题样式。

    @param document python-docx 文档对象。
    @return None。
    """
    config = _merge_format_config(config)
    normal = document.styles["Normal"]
    normal.font.name = config["ascii_font"]
    normal.font.size = Pt(config["body_size"])
    normal.font.color.rgb = TEXT_COLOR
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), config["body_font"])
    for style_name, size in (
        ("Title", config["heading1_size"] + 6),
        ("Heading 1", config["heading1_size"]),
        ("Heading 2", config["heading2_size"]),
        ("Heading 3", config["heading3_size"]),
    ):
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = config["ascii_font"]
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = TEXT_COLOR
            style._element.rPr.rFonts.set(qn("w:eastAsia"), config["heading_font"])


def _normalize_existing_paragraphs(document: Document, config: dict[str, Any] | None = None) -> None:
    """
    @brief 统一模板渲染后已有段落的字体和颜色。

    @param document python-docx 文档对象。
    @return None。

    docxtpl 渲染出的封面标题和生成时间来自模板段落，需要额外规范 run，
    避免 Word 主题样式造成中文字体或颜色不一致。
    """
    config = _merge_format_config(config)
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "Title":
            _apply_paragraph_font(
                paragraph,
                config["heading_font"],
                config["ascii_font"],
                size=config["heading1_size"] + 6,
                bold=True,
            )
        elif style_name.startswith("Heading"):
            _apply_paragraph_font(
                paragraph,
                config["heading_font"],
                config["ascii_font"],
                size=config["heading1_size"],
                bold=True,
            )
        else:
            _apply_paragraph_font(paragraph, config["body_font"], config["ascii_font"], size=config["body_size"])


def _apply_paragraph_font(
    paragraph,
    east_asia_font: str,
    ascii_font: str = ASCII_FONT,
    size: float = 11,
    bold: bool = False,
    first_line_indent_chars: float | None = None,
    line_spacing: float | None = None,
    space_before: Any | None = None,
    space_after: Any | None = None,
) -> None:
    """
    @brief 为段落中的所有 run 设置中英文字体。

    @param paragraph 段落对象。
    @param east_asia_font 中文字体。
    @param size 字号，单位为 pt。
    @param bold 是否加粗。
    @return None。
    """
    for run in paragraph.runs:
        run.font.name = ascii_font
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = TEXT_COLOR
        run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    if first_line_indent_chars is not None:
        paragraph.paragraph_format.first_line_indent = Pt(size * first_line_indent_chars)
    if line_spacing is not None:
        paragraph.paragraph_format.line_spacing = line_spacing
    if space_before is not None:
        paragraph.paragraph_format.space_before = space_before
    if space_after is not None:
        paragraph.paragraph_format.space_after = space_after


def _apply_cell_font(cell, config: dict[str, Any], bold: bool = False) -> None:
    """
    @brief 设置表格单元格字体。

    @param cell 表格单元格。
    @param east_asia_font 中文字体。
    @param bold 是否加粗。
    @return None。
    """
    for paragraph in cell.paragraphs:
        _apply_paragraph_font(
            paragraph,
            config["body_font"],
            config["ascii_font"],
            size=max(9, config["body_size"] - 0.5),
            bold=bold,
            line_spacing=1.15,
            space_after=Pt(0),
        )


def _merge_format_config(config: dict[str, Any] | None) -> dict[str, Any]:
    """
    @brief 合并并规范化 Word 格式配置。

    @param config 外部格式配置。
    @return 完整格式配置。
    """
    merged = dict(DEFAULT_FORMAT_CONFIG)
    if config:
        merged.update(config)
    for key in ("body_size", "heading1_size", "heading2_size", "heading3_size", "paragraph_space_after"):
        merged[key] = _as_float(merged.get(key), DEFAULT_FORMAT_CONFIG[key])
    if "body_space_after" not in merged and "paragraph_space_after" in merged:
        merged["body_space_after"] = {"value": merged["paragraph_space_after"], "unit": "pt"}
    for key in (
        "body_space_before",
        "body_space_after",
        "heading1_space_before",
        "heading1_space_after",
        "heading2_space_before",
        "heading2_space_after",
        "heading3_space_before",
        "heading3_space_after",
    ):
        merged[key] = _normalize_spacing_setting(merged.get(key), DEFAULT_FORMAT_CONFIG[key])
    merged["first_line_indent_chars"] = _as_float(merged.get("first_line_indent_chars"), 2)
    merged["line_spacing"] = _as_float(merged.get("line_spacing"), 1.5)
    merged["line_spacing_rule"] = _normalize_line_spacing_rule(merged.get("line_spacing_rule"), merged["line_spacing"])
    for key in ("body_font", "heading_font", "ascii_font", "table_style", "heading_numbering", "cover_style"):
        merged[key] = str(merged.get(key) or DEFAULT_FORMAT_CONFIG[key])
    merged["show_title"] = bool(merged.get("show_title", DEFAULT_FORMAT_CONFIG["show_title"]))
    return merged


def _as_float(value: Any, default: float) -> float:
    """
    @brief 将值转换为浮点数。

    @param value 原始值。
    @param default 默认值。
    @return 浮点数。
    """
    try:
        return float(value)
    except (TypeError, ValueError):
        return float(default)


def _normalize_line_spacing_rule(value: Any, default_multiple: float = 1.5) -> dict[str, Any]:
    """
    @brief 规范行距规则。

    @param value 原始行距规则。
    @param default_multiple 默认倍数行距。
    @return 规范后的行距规则。
    """
    if isinstance(value, dict):
        rule_type = str(value.get("type") or "multiple").lower()
        number = _as_float(value.get("value"), default_multiple)
        unit = str(value.get("unit") or ("pt" if rule_type == "exact" else "line")).lower()
    else:
        rule_type = "multiple"
        number = default_multiple
        unit = "line"
    if rule_type not in ("multiple", "exact"):
        rule_type = "multiple"
    if unit in ("磅", "point", "points"):
        unit = "pt"
    if unit in ("行", "lines"):
        unit = "line"
    if rule_type == "exact":
        unit = "pt"
    return {"type": rule_type, "value": max(0, number), "unit": unit}


def _resolve_line_spacing(config: dict[str, Any]) -> Any:
    """
    @brief 将行距规则转换为 python-docx 可用值。

    @param config 格式配置。
    @return 倍数或 Length 对象。
    """
    rule = config.get("line_spacing_rule") or {}
    if rule.get("type") == "exact":
        return Pt(_as_float(rule.get("value"), 23))
    return _as_float(rule.get("value"), config.get("line_spacing", 1.5))


def _normalize_spacing_setting(value: Any, default: dict[str, Any]) -> dict[str, Any]:
    """
    @brief 规范段前段后间距配置。

    @param value 原始配置，可为数字或 {"value": 数字, "unit": "pt/line"}。
    @param default 默认配置。
    @return 规范后的间距配置。
    """
    if isinstance(value, dict):
        number = _as_float(value.get("value"), default["value"])
        unit = str(value.get("unit") or default["unit"]).lower()
    else:
        number = _as_float(value, default["value"])
        unit = str(default["unit"]).lower()
    if unit in ("磅", "point", "points"):
        unit = "pt"
    if unit in ("行", "lines"):
        unit = "line"
    if unit not in ("pt", "line"):
        unit = default["unit"]
    return {"value": max(0, number), "unit": unit}


def _resolve_spacing(setting: dict[str, Any], font_size: float) -> Any:
    """
    @brief 将 pt/line 间距配置转换为 python-docx 可用长度。

    @param setting 间距配置。
    @param font_size 当前段落字号，单位 pt。
    @return python-docx Length 对象。
    """
    value = _as_float(setting.get("value"), 0)
    unit = str(setting.get("unit") or "pt").lower()
    if unit == "line":
        return Pt(value * font_size)
    return Pt(value)


def _heading_spacing(level: int, config: dict[str, Any], edge: str, font_size: float) -> Any:
    """
    @brief 获取指定标题层级的段前或段后间距。

    @param level 标题层级。
    @param config 格式配置。
    @param edge before 或 after。
    @param font_size 当前标题字号。
    @return python-docx Length 对象。
    """
    normalized_level = min(max(level, 1), 3)
    return _resolve_spacing(config[f"heading{normalized_level}_space_{edge}"], font_size)


def _heading_size(level: int, config: dict[str, Any]) -> float:
    """
    @brief 根据标题层级获取字号。

    @param level 标题层级。
    @param config 格式配置。
    @return 字号。
    """
    if level <= 1:
        return config["heading1_size"]
    if level == 2:
        return config["heading2_size"]
    return config["heading3_size"]


def _format_heading(heading: str, level: int, path: list[int], config: dict[str, Any]) -> str:
    """
    @brief 根据编号样式格式化标题文本。

    @param heading 原始标题。
    @param level 标题层级。
    @param path 章节序号路径。
    @param config 格式配置。
    @return 格式化标题。
    """
    title = _strip_heading_prefix(heading)
    if config.get("heading_numbering") == "chinese":
        prefix = _chinese_heading_prefix(level, path)
    else:
        prefix = ".".join(str(item) for item in path) + " "
    return f"{prefix}{title}"


def _strip_heading_prefix(heading: str) -> str:
    """
    @brief 移除已有标题编号。

    @param heading 原始标题。
    @return 去掉编号后的标题。
    """
    import re

    return re.sub(r"^(\d+(?:\.\d+)*[.、]?\s*|[一二三四五六七八九十]+、\s*|（[一二三四五六七八九十]+）\s*)", "", heading).strip()


def _chinese_heading_prefix(level: int, path: list[int]) -> str:
    """
    @brief 生成中文混合编号前缀。

    @param level 标题层级。
    @param path 章节序号路径。
    @return 编号前缀。
    """
    if level <= 1:
        return f"{_to_chinese_number(path[0])}、"
    if level == 2:
        return f"（{_to_chinese_number(path[-1])}）"
    return f"{path[-1]}. "


def _to_chinese_number(number: int) -> str:
    """
    @brief 将 1 到 99 的数字转换为中文序号。

    @param number 数字。
    @return 中文数字。
    """
    digits = "零一二三四五六七八九"
    if number <= 10:
        return "十" if number == 10 else digits[number]
    if number < 20:
        return f"十{digits[number % 10]}"
    tens, ones = divmod(number, 10)
    return f"{digits[tens]}十{digits[ones] if ones else ''}"


def _save_document_record(title: str, tree: dict[str, Any], output_path: Path) -> None:
    """
    @brief 保存生成文档记录。

    @param title 文档标题。
    @param tree 文档结构树。
    @param output_path 生成文件路径。
    @return None。
    """
    with get_connection() as connection:
        connection.execute(
            """
            INSERT INTO documents (title, topic, tree_json, docx_path, created_at, created_by)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                title,
                tree.get("title", title),
                json.dumps(tree, ensure_ascii=False),
                str(output_path),
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "anonymous",
            ),
        )
        connection.commit()
